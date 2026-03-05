#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
CV DOCX Generator — v2.0
Generates TWO professional CV variants from cv_data.json:
  - CV_ATS.docx   : ATS-optimized (1 column, clean text, keyword metadata)
  - CV_Visual.docx: Human-readable (styled headers, colors, hyperlinks)
"""

import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import re

# ─────────────────────────────────────────────
# ATS KEYWORDS injected into document metadata
# (safe, legitimate SEO — same as HTML meta tags)
# ─────────────────────────────────────────────
ATS_KEYWORDS = (
    "Ingeniero Mecatrónico, Automatización Industrial, Visión Artificial, "
    "Robótica, Python, OpenCV, YOLO, PLC Siemens, ROS, ESP32, Raspberry Pi, "
    "Machine Learning, Computer Vision, Control de Procesos, SCADA, HMI, "
    "Industria 4.0, IoT, Bogotá, Colombia, Remoto"
)


# ─────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────

def set_doc_metadata(doc, data, keywords: str):
    """Inject author, description and keyword metadata into DOCX core properties."""
    personal = data.get("informacion_personal", {})
    cp = doc.core_properties
    cp.author = personal.get("nombre", "")
    cp.title = "Curriculum Vitae — " + personal.get("nombre", "")
    cp.subject = "Ingeniero Mecatrónico | Automatización | Robótica | Visión Artificial"
    cp.keywords = keywords
    cp.description = data.get("resumen_profesional", "")[:255]
    cp.language = "es-CO"


def set_margins(doc, top=0.5, bottom=0.5, left=0.75, right=0.75):
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)


def add_hyperlink(paragraph, url: str, text: str, color: RGBColor = None):
    """Add a real clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)

    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
        rPr.append(c)

    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    hyperlink.append(r)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_hr(doc, color_hex="003366", thickness="12"):
    """Add a horizontal rule (bottom border on a paragraph)."""
    para = doc.add_paragraph()
    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), thickness)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    # Remove spacing
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(4)
    return para


def spacer(doc, size_pt=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(size_pt)


# ─────────────────────────────────────────────
# VISUAL CV SECTIONS
# ─────────────────────────────────────────────

def visual_section_title(doc, title: str):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(title.upper())
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    add_hr(doc, color_hex="003366", thickness="8")


def visual_header(doc, data):
    personal = data.get("informacion_personal", {})

    # Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(2)
    run = name_para.add_run(personal.get("nombre", "").upper())
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    # Contact line (plain text — ATS safe)
    phone = personal.get("telefono", "")
    email = personal.get("email", "")
    city = personal.get("ciudad", "")
    contact_text = "  |  ".join(filter(None, [phone, email, city]))
    contact_para = doc.add_paragraph(contact_text)
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.paragraph_format.space_before = Pt(0)
    contact_para.paragraph_format.space_after = Pt(2)
    contact_para.runs[0].font.size = Pt(9)
    contact_para.runs[0].font.color.rgb = RGBColor(60, 60, 60)

    # Links row with real hyperlinks
    links_para = doc.add_paragraph()
    links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    links_para.paragraph_format.space_before = Pt(0)
    links_para.paragraph_format.space_after = Pt(6)

    links = [
        ("LinkedIn", personal.get("linkedin", "")),
        ("GitHub", personal.get("github", "")),
        ("Portafolio", personal.get("portafolio", "")),
    ]
    first = True
    for label, url in links:
        if url:
            if not first:
                sep = links_para.add_run("  |  ")
                sep.font.size = Pt(9)
                sep.font.color.rgb = RGBColor(120, 120, 120)
            add_hyperlink(links_para, url, f"{label}: {url}", color=RGBColor(0, 80, 160))
            first = False

    add_hr(doc, "003366", "16")


def visual_summary(doc, data):
    summary = data.get("resumen_profesional", "")
    if not summary:
        return
    visual_section_title(doc, "Perfil Profesional")
    p = doc.add_paragraph(summary)
    p.runs[0].font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)


def visual_experience(doc, data):
    exps = data.get("experiencia", [])
    if not exps:
        return
    visual_section_title(doc, "Experiencia Profesional")
    for exp in exps:
        # Title | Company
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(1)
        r1 = p.add_run(exp.get("cargo", ""))
        r1.font.bold = True
        r1.font.size = Pt(11)
        r1.font.color.rgb = RGBColor(0, 51, 102)
        empresa = exp.get("empresa", "")
        if empresa:
            r2 = p.add_run(f"  —  {empresa}")
            r2.font.size = Pt(10)

        # Dates | Location
        fecha = f"{exp.get('fecha_inicio','')} – {exp.get('fecha_fin','Presente')}"
        ubicacion = exp.get("ubicacion", "")
        meta = "  |  ".join(filter(None, [fecha, ubicacion]))
        meta_p = doc.add_paragraph(meta)
        meta_p.runs[0].font.size = Pt(9)
        meta_p.runs[0].font.italic = True
        meta_p.runs[0].font.color.rgb = RGBColor(90, 90, 90)
        meta_p.paragraph_format.space_before = Pt(0)
        meta_p.paragraph_format.space_after = Pt(2)

        # Bullets
        for resp in exp.get("responsabilidades", []):
            if resp:
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.space_before = Pt(0)
                bp.paragraph_format.space_after = Pt(1)
                bp.add_run(resp).font.size = Pt(10)


def visual_education(doc, data):
    edus = data.get("educacion", [])
    if not edus:
        return
    visual_section_title(doc, "Educación")
    for edu in edus:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(1)
        r1 = p.add_run(edu.get("titulo", ""))
        r1.font.bold = True
        r1.font.size = Pt(11)
        r1.font.color.rgb = RGBColor(0, 51, 102)
        inst = edu.get("institucion", "")
        if inst:
            p.add_run(f"  —  {inst}").font.size = Pt(10)

        fecha = f"{edu.get('fecha_inicio','')} – {edu.get('fecha_fin','')}"
        ubicacion = edu.get("ubicacion", "")
        meta = "  |  ".join(filter(None, [fecha, ubicacion]))
        mp = doc.add_paragraph(meta)
        mp.runs[0].font.size = Pt(9)
        mp.runs[0].font.italic = True
        mp.paragraph_format.space_before = Pt(0)
        mp.paragraph_format.space_after = Pt(1)

        gpa = edu.get("gpa", "")
        dist = edu.get("distinciones", "")
        if gpa:
            gp = doc.add_paragraph(f"GPA: {gpa}")
            gp.runs[0].font.size = Pt(9)
            gp.paragraph_format.space_before = Pt(0)
            gp.paragraph_format.space_after = Pt(1)
        if dist:
            dp = doc.add_paragraph(f"• {dist}")
            dp.runs[0].font.size = Pt(9)
            dp.paragraph_format.space_before = Pt(0)
            dp.paragraph_format.space_after = Pt(1)


def visual_projects(doc, data):
    projects = data.get("proyectos", [])
    if not projects:
        return
    visual_section_title(doc, "Proyectos Destacados")
    for proj in projects:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(1)
        r1 = p.add_run(proj.get("nombre", ""))
        r1.font.bold = True
        r1.font.size = Pt(11)
        r1.font.color.rgb = RGBColor(0, 51, 102)
        fecha = proj.get("fecha", "")
        if fecha:
            p.add_run(f"  ({fecha})").font.size = Pt(10)

        desc = proj.get("descripcion", "")
        if desc:
            dp = doc.add_paragraph(desc)
            dp.runs[0].font.size = Pt(10)
            dp.paragraph_format.space_before = Pt(1)
            dp.paragraph_format.space_after = Pt(2)

        techs = proj.get("tecnologias", [])
        if techs:
            tech_str = " · ".join(techs) if isinstance(techs, list) else techs
            tp = doc.add_paragraph()
            tr = tp.add_run("Tecnologías: ")
            tr.font.bold = True
            tr.font.size = Pt(9)
            tp.add_run(tech_str).font.size = Pt(9)
            tp.paragraph_format.space_before = Pt(0)
            tp.paragraph_format.space_after = Pt(1)

        resultados = proj.get("resultados", [])
        if resultados:
            if isinstance(resultados, list):
                for res in resultados:
                    rp = doc.add_paragraph(f"✓ {res}")
                    rp.runs[0].font.size = Pt(9)
                    rp.runs[0].font.color.rgb = RGBColor(0, 110, 0)
                    rp.paragraph_format.space_before = Pt(0)
                    rp.paragraph_format.space_after = Pt(1)
            else:
                rp = doc.add_paragraph(f"✓ {resultados}")
                rp.runs[0].font.size = Pt(9)
                rp.paragraph_format.space_before = Pt(0)
                rp.paragraph_format.space_after = Pt(1)


def visual_skills(doc, data):
    skills = data.get("habilidades", {})
    if not skills:
        return
    visual_section_title(doc, "Habilidades Técnicas")

    labels = {
        "lenguajes_programacion": "Programación",
        "software_diseno": "CAD / Diseño",
        "robotica_automatizacion": "Robótica y Automatización",
        "vision_artificial": "Visión Artificial / IA",
        "hardware_electronica": "Hardware y Electrónica",
        "desarrollo_software": "Desarrollo de Software",
        "manufactura": "Manufactura",
    }
    for key, label in labels.items():
        items = skills.get(key, {})
        if not items:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(f"{label}: ")
        r.font.bold = True
        r.font.size = Pt(10)
        text = ", ".join(
            f"{k} ({v})" if v else k for k, v in items.items()
        )
        p.add_run(text).font.size = Pt(10)


def visual_languages(doc, data):
    langs = data.get("idiomas", {})
    if not langs:
        return
    visual_section_title(doc, "Idiomas")
    p = doc.add_paragraph("  |  ".join(f"{k}: {v}" for k, v in langs.items()))
    p.runs[0].font.size = Pt(10)


# ─────────────────────────────────────────────
# ATS CV (plain text, no colors, no icons)
# ─────────────────────────────────────────────

def ats_section_title(doc, title: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title.upper())
    r.font.size = Pt(12)
    r.font.bold = True
    add_hr(doc, "000000", "8")


def ats_header(doc, data):
    personal = data.get("informacion_personal", {})
    name_p = doc.add_paragraph(personal.get("nombre", "").upper())
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.runs[0].font.size = Pt(18)
    name_p.runs[0].font.bold = True
    name_p.paragraph_format.space_after = Pt(2)

    phone = personal.get("telefono", "")
    email = personal.get("email", "")
    city = personal.get("ciudad", "")
    contact_text = " | ".join(filter(None, [phone, email, city]))
    cp = doc.add_paragraph(contact_text)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.runs[0].font.size = Pt(10)
    cp.paragraph_format.space_before = Pt(0)
    cp.paragraph_format.space_after = Pt(2)

    links = []
    if personal.get("linkedin"):
        links.append(f"LinkedIn: {personal['linkedin']}")
    if personal.get("github"):
        links.append(f"GitHub: {personal['github']}")
    if personal.get("portafolio"):
        links.append(f"Portafolio: {personal['portafolio']}")

    if links:
        lp = doc.add_paragraph(" | ".join(links))
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.runs[0].font.size = Pt(9)
        lp.paragraph_format.space_before = Pt(0)
        lp.paragraph_format.space_after = Pt(4)

    add_hr(doc, "000000", "12")


def ats_summary(doc, data):
    summary = data.get("resumen_profesional", "")
    if not summary:
        return
    ats_section_title(doc, "Perfil Profesional")
    p = doc.add_paragraph(summary)
    p.runs[0].font.size = Pt(10)


def ats_experience(doc, data):
    exps = data.get("experiencia", [])
    if not exps:
        return
    ats_section_title(doc, "Experiencia Profesional")
    for exp in exps:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        r = p.add_run(exp.get("cargo", ""))
        r.font.bold = True
        r.font.size = Pt(11)
        empresa = exp.get("empresa", "")
        if empresa:
            p.add_run(f" — {empresa}").font.size = Pt(10)

        fecha = f"{exp.get('fecha_inicio','')} - {exp.get('fecha_fin','Presente')}"
        ubicacion = exp.get("ubicacion", "")
        meta = " | ".join(filter(None, [fecha, ubicacion]))
        mp = doc.add_paragraph(meta)
        mp.runs[0].font.size = Pt(9)
        mp.paragraph_format.space_before = Pt(0)
        mp.paragraph_format.space_after = Pt(2)

        for resp in exp.get("responsabilidades", []):
            if resp:
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.space_before = Pt(0)
                bp.paragraph_format.space_after = Pt(1)
                bp.add_run(resp).font.size = Pt(10)


def ats_education(doc, data):
    edus = data.get("educacion", [])
    if not edus:
        return
    ats_section_title(doc, "Educación")
    for edu in edus:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        r = p.add_run(edu.get("titulo", ""))
        r.font.bold = True
        r.font.size = Pt(11)
        if edu.get("institucion"):
            p.add_run(f" — {edu['institucion']}").font.size = Pt(10)

        fecha = f"{edu.get('fecha_inicio','')} - {edu.get('fecha_fin','')}"
        ubicacion = edu.get("ubicacion", "")
        meta = " | ".join(filter(None, [fecha, ubicacion]))
        mp = doc.add_paragraph(meta)
        mp.runs[0].font.size = Pt(9)
        mp.paragraph_format.space_before = Pt(0)
        mp.paragraph_format.space_after = Pt(1)

        if edu.get("gpa"):
            gp = doc.add_paragraph(f"GPA: {edu['gpa']}")
            gp.runs[0].font.size = Pt(9)
            gp.paragraph_format.space_before = Pt(0)
        if edu.get("distinciones"):
            dp = doc.add_paragraph(f"- {edu['distinciones']}")
            dp.runs[0].font.size = Pt(9)
            dp.paragraph_format.space_before = Pt(0)


def ats_projects(doc, data):
    projects = data.get("proyectos", [])
    if not projects:
        return
    ats_section_title(doc, "Proyectos Destacados")
    for proj in projects:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        r = p.add_run(proj.get("nombre", ""))
        r.font.bold = True
        r.font.size = Pt(11)
        fecha = proj.get("fecha", "")
        if fecha:
            p.add_run(f" ({fecha})").font.size = Pt(10)

        desc = proj.get("descripcion", "")
        if desc:
            dp = doc.add_paragraph(desc)
            dp.runs[0].font.size = Pt(10)
            dp.paragraph_format.space_before = Pt(1)

        techs = proj.get("tecnologias", [])
        if techs:
            tech_str = " | ".join(techs) if isinstance(techs, list) else techs
            tp = doc.add_paragraph()
            tp.add_run("Tecnologias: ").font.bold = True
            tp.add_run(tech_str).font.size = Pt(9)
            tp.paragraph_format.space_before = Pt(0)

        resultados = proj.get("resultados", [])
        if resultados:
            items = resultados if isinstance(resultados, list) else [resultados]
            for res in items:
                rp = doc.add_paragraph(f"- {res}")
                rp.runs[0].font.size = Pt(9)
                rp.paragraph_format.space_before = Pt(0)
                rp.paragraph_format.space_after = Pt(1)


def ats_skills(doc, data):
    skills = data.get("habilidades", {})
    if not skills:
        return
    ats_section_title(doc, "Habilidades Tecnicas")

    labels = {
        "lenguajes_programacion": "Programacion",
        "software_diseno": "Software CAD",
        "robotica_automatizacion": "Robotica y Automatizacion",
        "vision_artificial": "Vision Artificial e IA",
        "hardware_electronica": "Hardware y Electronica",
        "desarrollo_software": "Herramientas de Desarrollo",
        "manufactura": "Manufactura",
    }
    for key, label in labels.items():
        items = skills.get(key, {})
        if not items:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(1)
        p.add_run(f"{label}: ").font.bold = True
        text = ", ".join(
            f"{k} ({v})" if v else k for k, v in items.items()
        )
        p.add_run(text).font.size = Pt(10)


def ats_languages(doc, data):
    langs = data.get("idiomas", {})
    if not langs:
        return
    ats_section_title(doc, "Idiomas")
    p = doc.add_paragraph(" | ".join(f"{k}: {v}" for k, v in langs.items()))
    p.runs[0].font.size = Pt(10)


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────

def build_visual(data) -> Document:
    doc = Document()
    set_margins(doc, 0.5, 0.5, 0.75, 0.75)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    set_doc_metadata(doc, data, ATS_KEYWORDS)

    visual_header(doc, data)
    visual_summary(doc, data)
    visual_experience(doc, data)
    visual_education(doc, data)
    visual_projects(doc, data)
    visual_skills(doc, data)
    visual_languages(doc, data)
    return doc


def build_ats(data) -> Document:
    doc = Document()
    set_margins(doc, 0.6, 0.6, 1.0, 1.0)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    set_doc_metadata(doc, data, ATS_KEYWORDS)

    ats_header(doc, data)
    ats_summary(doc, data)
    ats_experience(doc, data)
    ats_education(doc, data)
    ats_projects(doc, data)
    ats_skills(doc, data)
    ats_languages(doc, data)
    return doc


def main():
    print("=" * 55)
    print("   CV DOCX Generator v2.0")
    print("=" * 55)

    json_path = Path("cv_data.json")
    if not json_path.exists():
        print("[ERROR] No se encontro cv_data.json en el directorio actual.")
        return

    with open(json_path, "r", encoding="utf-8") as f:
        cv_data = json.load(f)

    nombre = cv_data.get("informacion_personal", {}).get("nombre", "CV").split()[0]

    # Visual version
    print("\n[1/2] Generando CV_Visual.docx ...")
    doc_visual = build_visual(cv_data)
    out_visual = Path(f"CV_Visual_{nombre}.docx")
    doc_visual.save(out_visual)
    print(f"      Guardado: {out_visual}")

    # ATS version
    print("[2/2] Generando CV_ATS.docx ...")
    doc_ats = build_ats(cv_data)
    out_ats = Path(f"CV_ATS_{nombre}.docx")
    doc_ats.save(out_ats)
    print(f"      Guardado: {out_ats}")

    print("\n" + "=" * 55)
    print(f"  Listo! Dos archivos generados:")
    print(f"  - {out_visual}  -->  Envia directo a reclutadores humanos")
    print(f"  - {out_ats}     -->  Sube a Computrabajo, LinkedIn, ATS")
    print("=" * 55)
    print("\nMetadatos inyectados:")
    print(f"  Keywords: {ATS_KEYWORDS[:80]}...")
    print("\nSiguiente paso: python bot/bot.py --dry-run")


if __name__ == "__main__":
    main()
